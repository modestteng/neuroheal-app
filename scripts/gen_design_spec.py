"""
生成软件著作权登记申请用的《软件设计说明书》Word 文档。

输出：软著材料/软件设计说明书.docx

格式要求（来自中国版权保护中心 2026 要求）：
- A4 Word，建议 20–30 页
- 页眉：软件全称 + 版本号（左），第 X 页 / 共 Y 页（右）
- 必须包含：操作系统 / 开发工具 / 硬件需求 / 编程语言
- 图文并茂，每个功能模块配截图（本脚本用占位框，用户后续补图）
- 禁忌：不含真实用户数据、不含医学诊断、不含营销语
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "软著材料" / "软件设计说明书.docx"
SHOT_DIR = ROOT / "软著材料" / "screenshots"

# 图编号 → 截图文件映射；架构图由 scripts/gen_diagrams.py 单独渲染
FIGURE_TO_SHOT: dict[str, str] = {
    "3.1.1": "10_arch_diagram.png",
    "3.2.1": "11_loop_diagram.png",
    "4.1.1": "01_home_main.png",
    "4.2.1": "03_train_main.png",
    "4.3.1": "08_community_main.png",
    "4.4.1": "04_sub_aichat.png",
    "4.5.1": "05_sub_player.png",
    "4.6.1": "06_sub_race.png",
    "4.7.1": "07_sub_prescription.png",
    "4.7.2": "02_sub_device.png",
    "8.1.1": "01_home_main.png",
    "8.2.1": "01_home_main.png",
    "8.3.1": "03_train_main.png",
    "8.4.1": "08_community_main.png",
    "8.5.1": "04_sub_aichat.png",
    "8.6.1": "05_sub_player.png",
    "8.7.1": "06_sub_race.png",
    "8.8.1": "07_sub_prescription.png",
}

SOFTWARE_FULL_NAME = "智愈莘莘 NeuroHeal 脑电情绪监测与干预反馈系统"
SOFTWARE_VERSION = "V1.0"
SOFTWARE_TITLE = f"{SOFTWARE_FULL_NAME} {SOFTWARE_VERSION}"

CN_FONT = "宋体"
EN_FONT = "Times New Roman"


def set_font_to_run(run, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.name = EN_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    # 中文字体单独设置
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)


def add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font_to_run(run, size_pt=sizes.get(level, 12), bold=True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font_to_run(run, size_pt=10.5)


def add_list_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font_to_run(run, size_pt=10.5)


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font_to_run(run, size_pt=10, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_font_to_run(run, size_pt=10)
    # 表后空一段
    doc.add_paragraph()


def _extract_fig_number(caption: str) -> str | None:
    """从图注 '图 4.1.1  xxx' 中提取 '4.1.1'。"""
    import re
    m = re.match(r"图\s+(\d+(?:\.\d+){1,2})", caption)
    return m.group(1) if m else None


def add_figure_placeholder(doc: Document, caption: str, height_cm: float = 7.5) -> None:
    """插入截图（如果有对应文件）或占位框 + 居中图注。"""
    fig_no = _extract_fig_number(caption)
    shot_path = None
    if fig_no and fig_no in FIGURE_TO_SHOT:
        candidate = SHOT_DIR / FIGURE_TO_SHOT[fig_no]
        if candidate.exists():
            shot_path = candidate

    if shot_path is not None:
        # 直接插入图片，居中。自动按图像纵横比选择宽度或高度约束。
        from PIL import Image as PILImage  # 仅这里用，避免顶部强依赖

        with PILImage.open(shot_path) as im:
            w, h = im.size

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(4)
        p_img.paragraph_format.space_after = Pt(2)
        p_img.paragraph_format.keep_with_next = True  # 保证图与下方图注不被分页
        run_img = p_img.add_run()
        if w > h * 1.3:
            # 横向图（架构图）：按宽度约束 15cm，正文区刚好
            run_img.add_picture(str(shot_path), width=Cm(15.0))
        else:
            # 纵向图（手机截图）：按高度约束 height_cm
            run_img.add_picture(str(shot_path), height=Cm(height_cm))
    else:
        # 保留占位框（用于尚未截图的架构图等）
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.rows[0].cells[0]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), str(int(height_cm * 567)))
        trHeight.set(qn("w:hRule"), "exact")
        trPr.append(trHeight)

        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        tcPr.append(shd)

        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "dashed")
            b.set(qn("w:sz"), "6")
            b.set(qn("w:color"), "999999")
            tcBorders.append(b)
        tcPr.append(tcBorders)

        cell.text = ""
        run = cell.paragraphs[0].add_run("【此处粘贴截图：" + (fig_no or "—") + "】")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_to_run(run, size_pt=11)
        run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    # 图注
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(10)
    cap_run = cap.add_run(caption)
    set_font_to_run(cap_run, size_pt=9.5)
    cap_run.italic = True


def set_a4_margins(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.0)


def _add_page_field(run) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_instr = OxmlElement("w:instrText")
    fld_instr.set(qn("xml:space"), "preserve")
    fld_instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(fld_instr)
    run._element.append(fld_end)


def _add_numpages_field(run) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_instr = OxmlElement("w:instrText")
    fld_instr.set(qn("xml:space"), "preserve")
    fld_instr.text = "NUMPAGES"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(fld_instr)
    run._element.append(fld_end)


def add_header(section) -> None:
    """页眉：用 2 列 1 行无边框表格——软件全称 + 版本号（左），第 X 页 / 共 Y 页（右）。"""
    header = section.header
    # 清空默认段落
    header.paragraphs[0].clear()

    table = header.add_table(rows=1, cols=2, width=Cm(16.3))
    table.autofit = False
    # 列宽
    table.columns[0].width = Cm(12.0)
    table.columns[1].width = Cm(4.3)

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    # 去除单元格四边边框
    for cell in (left_cell, right_cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "nil")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    # 左：软件全称
    left_p = left_cell.paragraphs[0]
    left_p.text = ""
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_l = left_p.add_run(SOFTWARE_TITLE)
    set_font_to_run(run_l, size_pt=9)

    # 右：第 X 页 / 共 Y 页
    right_p = right_cell.paragraphs[0]
    right_p.text = ""
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = right_p.add_run("第 "); set_font_to_run(r1, size_pt=9)
    r_page = right_p.add_run(); set_font_to_run(r_page, size_pt=9); _add_page_field(r_page)
    r2 = right_p.add_run(" 页 / 共 "); set_font_to_run(r2, size_pt=9)
    r_total = right_p.add_run(); set_font_to_run(r_total, size_pt=9); _add_numpages_field(r_total)
    r3 = right_p.add_run(" 页"); set_font_to_run(r3, size_pt=9)

    # 表格底部加细线
    tbl_pr = table._element.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tbl_borders.append(b)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "808080")
    tbl_borders.append(bottom)
    tbl_pr.append(tbl_borders)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_cover(doc: Document) -> None:
    """封面页：软件名 + 版本 + 副标题。"""
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SOFTWARE_FULL_NAME)
    set_font_to_run(run, size_pt=24, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(SOFTWARE_VERSION)
    set_font_to_run(run2, size_pt=18, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("软件设计说明书")
    set_font_to_run(run3, size_pt=22, bold=True)

    for _ in range(6):
        doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("文档版本：1.0")
    set_font_to_run(run4, size_pt=12)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("编写日期：2026 年 5 月")
    set_font_to_run(run5, size_pt=12)

    add_page_break(doc)


# ============ 章节内容 ============

def chapter_1(doc: Document) -> None:
    add_heading(doc, "第 1 章  引言", 1)

    add_heading(doc, "1.1 编写目的", 2)
    add_body(doc, (
        f"本文档是《{SOFTWARE_TITLE}》的软件设计说明书，用于完整描述本软件的系统目标、"
        "总体架构、主要功能模块、关键数据结构、对外接口和运行环境。本文档面向以下读者："
    ))
    add_list_item(doc, "软件开发与维护人员：理解系统模块划分、接口约定、数据流向；")
    add_list_item(doc, "软件评审与测试人员：核对功能实现是否符合设计预期；")
    add_list_item(doc, "知识产权审查人员：作为软件著作权登记的设计说明依据。")

    add_heading(doc, "1.2 项目背景", 2)
    add_body(doc, (
        "高校学生群体长期面临学业压力、人际适应、睡眠失衡等多重心理压力来源，"
        "传统的问卷量表与定期访谈难以反映学生情绪的实时波动，也无法在情绪低谷时刻"
        "立即提供可执行的干预建议。本系统基于便携式 EEG（脑电）信号采集设备，"
        "实时获取用户脑波信号，识别 Valence 情绪效价等级，并由 AI 反馈模块生成"
        "温和、可执行的安抚语言与训练推荐，形成可持续运转的反馈干预闭环。"
    ))
    add_body(doc, (
        "本系统为辅助干预性质的软件工具，不具有医学诊断或治疗效力，"
        "结果与建议仅作为用户日常心理健康自我管理的参考。"
    ))

    add_heading(doc, "1.3 术语与缩略语", 2)
    add_table(
        doc,
        header=["术语 / 缩略语", "中文含义", "本文档中的用途"],
        rows=[
            ["EEG", "脑电图 / 脑电信号", "采集设备输出的原始生理信号源"],
            ["Valence", "情绪效价", "用于描述情绪积极程度的连续量化指标"],
            ["AI", "人工智能", "本系统中指 DeepSeek 大语言模型反馈模块"],
            ["PWA", "渐进式 Web 应用", "前端发布形态，可安装、可离线"],
            ["React", "前端组件化框架", "前端 UI 主体框架（v19）"],
            ["TypeScript", "强类型 JavaScript 超集", "前端开发语言"],
            ["Vite", "现代前端构建工具", "用于开发、构建与打包"],
            ["API", "应用程序接口", "前端与后端代理之间的 HTTP 接口"],
        ],
    )

    add_heading(doc, "1.4 参考资料", 2)
    add_list_item(doc, "GB/T 8567-2006《计算机软件文档编制规范》")
    add_list_item(doc, "《计算机软件著作权登记办法》（国家版权局）")
    add_list_item(doc, "React 19 官方文档、Vite 8 官方文档、DeepSeek 平台 API 文档")


def chapter_2(doc: Document) -> None:
    add_heading(doc, "第 2 章  系统概述", 1)

    add_heading(doc, "2.1 系统定位", 2)
    add_body(doc, (
        f"《{SOFTWARE_TITLE}》是一款面向高校学生心理健康辅助场景的脑电情绪监测与"
        "AI 反馈干预闭环演示系统。系统通过便携式 EEG 设备持续采集用户脑波信号，"
        "识别 Valence 情绪效价等级，由 AI 模块生成温和的安抚语言与具体的下一步干预"
        "建议，并支持用户进入呼吸训练、数字处方、调节小游戏或 AI 心理陪伴等模块；"
        "干预完成后系统继续监测脑电信号、再次进行效价判断，形成"
        "「监测 → 判断 → 反馈 → 干预 → 再监测」的可持续闭环。"
    ))

    add_heading(doc, "2.2 主要功能模块", 2)
    add_table(
        doc,
        header=["模块名称", "对应入口", "主要职责"],
        rows=[
            ["监测模块", "首页 / Home", "EEG 实时波形显示、Valence 等级判定、AI 反馈卡片、推荐干预入口"],
            ["干预模块", "干预页 / Train", "干预动作列表、数字处方包导航、呼吸训练 / 调节游戏 / AI 陪伴入口"],
            ["记录模块", "记录页 / Community", "闭环时间轴、成长档案、勋章墙、本月打卡概览"],
            ["AI 反馈代理", "后端服务（/api/chat）", "接收前端 Valence 上下文，调用 DeepSeek 生成温和反馈"],
            ["设备管理", "子页 / Device", "脑电头环连接状态、电极信号质量、采样率与校准流程"],
            ["数字处方", "子页 / Prescription", "考研专注流、安睡曲、疗愈包、社交自信课 4 套主题处方"],
            ["呼吸训练", "子页 / Player", "4-7-8 呼吸节律引导、实时放松度反馈"],
            ["AI 心理陪伴", "子页 / AiChat", "与小愈的多轮对话、自动注入 Valence 上下文"],
            ["调节小游戏", "子页 / Race", "意念赛车，专注度驱动的轻量调节游戏"],
            ["成长档案", "子页 / Growth", "打卡日历、勋章墙、本周心理风险评估"],
        ],
    )

    add_heading(doc, "2.3 运行环境", 2)
    add_table(
        doc,
        header=["类别", "项目", "推荐 / 说明"],
        rows=[
            ["操作系统", "前端浏览器端", "Windows 10/11、macOS 12+、iOS 16+、Android 12+"],
            ["操作系统", "后端运行端", "Windows 10/11、Ubuntu 20.04+、macOS 12+"],
            ["运行时", "Node.js", "≥ 20.11.0（用于后端 DeepSeek 代理服务）"],
            ["浏览器", "前端渲染端", "Chrome 110+、Edge 110+、Safari 16+ 或同等内核"],
            ["硬件需求", "客户端", "1 GHz 双核 CPU、4 GB 内存、500 MB 可用存储、支持蓝牙 4.0+"],
            ["硬件需求", "服务端", "1 GHz 双核 CPU、512 MB 内存、500 MB 可用存储"],
            ["生理信号设备", "EEG 头环", "支持 256 Hz 采样的便携式干电极脑电头环（演示阶段使用 mock 数据）"],
            ["网络", "依赖", "HTTPS 出站访问 api.deepseek.com（仅 AI 反馈调用使用）"],
        ],
    )

    add_heading(doc, "2.4 开发环境与编程语言", 2)
    add_table(
        doc,
        header=["类别", "选型", "版本"],
        rows=[
            ["前端编程语言", "TypeScript", "6.x（编译目标 ES2023）"],
            ["前端框架", "React", "19.2"],
            ["前端构建工具", "Vite", "8.0"],
            ["UI 动画库", "Framer Motion", "12.x"],
            ["图表库", "Recharts", "3.x"],
            ["图标库", "Lucide React", "1.x"],
            ["后端语言", "JavaScript（ES Modules）", "Node.js 20.11 内置"],
            ["后端 HTTP", "Node 原生 http 模块", "—"],
            ["开发工具", "Visual Studio Code", "1.85+"],
            ["代码规范", "ESLint + typescript-eslint", "10.x / 8.x"],
            ["包管理", "npm", "10.x（随 Node 自带）"],
        ],
    )


def chapter_3(doc: Document) -> None:
    add_heading(doc, "第 3 章  总体设计", 1)

    add_heading(doc, "3.1 系统架构", 2)
    add_body(doc, (
        "本系统采用前后端分离的轻量架构，分为三层：用户交互层（前端 PWA）、"
        "AI 反馈代理层（Node.js 后端）、外部 AI 服务层（DeepSeek 大模型 API）。"
        "前端通过 Vite 开发服务器代理或生产环境同源访问，统一以 /api/* 路径访问后端，"
        "后端只负责承载 API 密钥并对调用进行轻量整形与安全校验，自身不存储用户数据。"
    ))
    add_figure_placeholder(doc, "图 3.1.1  系统总体架构示意图", height_cm=8.0)

    add_heading(doc, "3.2 核心反馈闭环", 2)
    add_body(doc, (
        "系统主干流程是一个可持续运转的闭环：「EEG 信号采集 → Valence 情绪效价识别 → "
        "AI 安抚 / 建议 / 训练推荐 → 用户执行干预 → 再次采集 EEG → 再次进行 Valence 判断」。"
        "前端通过 ValenceLoopProvider 全局状态钩子驱动该闭环：定时器每 9 秒推进一帧"
        "Valence 快照，触发 UI 更新；当用户主动点击「请求 AI 反馈」或进入 AI 陪伴页时，"
        "系统将当前 Valence 等级、置信度与 EEG 摘要作为隐藏的系统上下文一并提交，"
        "后端再将其与「小愈」系统提示词合并后转发至 DeepSeek，得到针对当前状态的"
        "具体反馈语句与下一步建议。"
    ))
    add_figure_placeholder(doc, "图 3.2.1  Valence 闭环时序与数据流", height_cm=8.0)

    add_heading(doc, "3.3 模块划分与依赖关系", 2)
    add_table(
        doc,
        header=["层", "模块", "说明"],
        rows=[
            ["展示层", "Screens / 主页面", "Home / Train / Community 三个主标签页"],
            ["展示层", "Screens / 子页面", "Device / Prescription / Player / AiChat / Growth / Race 共 6 个"],
            ["展示层", "Components", "PhoneFrame、BottomNav、SubScreen、UI 工具组件库"],
            ["状态层", "useValenceLoop", "Valence 闭环驱动器（定时器、AI 调用、历史维护）"],
            ["状态层", "nav", "页面路由与子页堆栈管理 Context"],
            ["状态层", "game-session", "意念赛车的实时专注度会话状态"],
            ["数据层", "data/mock", "演示数据：Valence 快照池、干预动作、闭环记录、勋章等"],
            ["服务层", "server/deepseek.mjs", "Node.js 后端 AI 反馈代理（仅 /api/chat、/api/health）"],
        ],
    )

    add_heading(doc, "3.4 设计原则", 2)
    add_list_item(doc, "主线优先：界面信息层级以「监测 / 干预 / 记录」三主线为骨架，辅助功能不抢占首屏；")
    add_list_item(doc, "实时反馈：状态变更必须能在 10 秒内驱动 UI 更新，避免阻塞用户感知；")
    add_list_item(doc, "温和陪伴：所有 AI 反馈文案以倾听、安抚、可执行建议为主，不出现医学诊断；")
    add_list_item(doc, "辅助性界定：所有结果与建议明确标注为辅助参考，提供线下专业资源入口；")
    add_list_item(doc, "可演示性：演示数据完全 mock 化，不依赖真实硬件即可运行完整闭环。")


def chapter_4(doc: Document) -> None:
    add_heading(doc, "第 4 章  详细功能设计", 1)

    add_heading(doc, "4.1 监测模块（首页 / Home）", 2)
    add_body(doc, (
        "监测模块是系统的入口和主控制台。首屏自上而下依次为：实时 EEG-Valence 控制台卡、"
        "闭环运行流程示意条、AI 当前反馈与推荐干预卡。Valence 控制台显示当前情绪效价"
        "等级、分数（0-100）、置信度与信号质量；实时 EEG 波形以 28 点采样图展示，"
        "振幅会根据当前 Valence 等级动态调整（低效价 → 大幅波动、高效价 → 平稳）。"
    ))
    add_body(doc, (
        "AI 反馈卡片显示模型针对当前状态生成的安抚文字；用户可点击「请求 AI 反馈」按钮"
        "主动触发一次新的 DeepSeek 调用，按钮处于 loading 状态时显示旋转图标。"
        "推荐干预区根据当前 Valence 快照中的 recommendedActionId 字段映射到具体干预动作，"
        "点击「开始」按钮可直接跳转到对应的子页面（呼吸训练 / 数字处方 / 调节游戏 / AI 陪伴）。"
    ))
    add_table(
        doc,
        header=["步骤", "用户操作", "系统响应"],
        rows=[
            ["1", "打开应用，进入首页", "ValenceLoopProvider 自动启动，每 9 秒推进一帧 Valence 快照"],
            ["2", "观察 EEG 波形与 Valence 分数", "GaugeRing 数字滚动至当前分数，等级标签同步刷新"],
            ["3", "点击「请求 AI 反馈」按钮", "前端拼装 system + user 消息，POST /api/chat"],
            ["4", "等待响应（典型 2–4 秒）", "按钮显示 loading；反馈卡片渐显新文案"],
            ["5", "点击推荐干预的「开始」按钮", "跳转到对应子页（player / race / aichat / prescription）"],
            ["6", "返回首页", "tick 持续运行，Valence 等级根据新一帧快照继续更新"],
        ],
    )
    add_figure_placeholder(doc, "图 4.1.1  监测首页 · EEG-Valence 控制台与 AI 反馈卡", height_cm=10.0)

    add_heading(doc, "4.2 干预模块（干预页 / Train）", 2)
    add_body(doc, (
        "干预模块集中展示所有可选的干预方式。页面顶部展示当前状态与 AI 推荐的优先动作，"
        "中部依次列出 4 种干预动作（呼吸训练、调节游戏、AI 陪伴、数字处方），"
        "下部展示 4 套数字处方包（考研冲刺专注流、失眠改善安睡曲、情感创伤疗愈包、社交自信提升课），"
        "每套处方均包含适用人群、推荐脑波频段、单次时长与节数说明。"
    ))
    add_figure_placeholder(doc, "图 4.2.1  干预页 · AI 推荐动作与数字处方列表", height_cm=10.0)

    add_heading(doc, "4.3 记录模块（记录页 / Community）", 2)
    add_body(doc, (
        "记录模块用于归档用户的闭环干预历史。页面包含四个区块：当前 Valence 状态卡（含连续打卡天数与本月打卡进度条）、"
        "闭环记录时间轴（每条记录展示干预前等级、行动、干预后等级、Δ score 增量）、"
        "成长档案三宫格（连续打卡 / 本月打卡 / 已获勋章）、以及已获勋章条带。"
    ))
    add_figure_placeholder(doc, "图 4.3.1  记录页 · 闭环时间轴与成长档案", height_cm=10.0)

    add_heading(doc, "4.4 AI 心理陪伴（子页 / AiChat）", 2)
    add_body(doc, (
        "AI 心理陪伴模块以聊天界面与用户进行多轮温和对话。每次发送消息时，"
        "前端自动在 messages 数组最前面插入一条隐藏的 system 类型消息，携带当前 Valence "
        "等级、score、置信度、信号质量与 EEG 摘要，使 DeepSeek 模型能基于实时状态调整"
        "回复语气与建议方向。界面顶部展示当前 Valence 状态条，让用户清楚 AI 「知道」什么。"
        "页面同时提供「预约校心理中心」与「24h 援助热线」两个快捷入口。"
    ))
    add_figure_placeholder(doc, "图 4.4.1  AI 陪伴子页 · 实时 Valence 状态与对话流", height_cm=10.0)

    add_heading(doc, "4.5 呼吸训练（子页 / Player）", 2)
    add_body(doc, (
        "呼吸训练子页基于 4-7-8 呼吸法实现。中央动态光圈随呼吸节奏放大与收缩："
        "吸气阶段 4 秒、屏息 7 秒、呼气 8 秒，循环往复。播放器底部显示已练习时长、"
        "实时放松度估计值与总时长。训练完成后弹出结果卡，展示放松度提升量并写入成长档案。"
    ))
    add_figure_placeholder(doc, "图 4.5.1  呼吸训练子页 · 4-7-8 节律引导界面", height_cm=10.0)

    add_heading(doc, "4.6 意念赛车（子页 / Race）", 2)
    add_body(doc, (
        "意念赛车是面向「专注度调节」的轻量游戏化干预模块。用户按住屏幕中央的按钮持续输出"
        "专注力，赛车前进速度与按压时长成正比；松开则减速。游戏分为 Ready、倒计时、Running、"
        "Summary 四个阶段，全程 30 秒。完成后系统继续监测 EEG 信号，更新 Valence 结果，回到闭环主线。"
    ))
    add_figure_placeholder(doc, "图 4.6.1  意念赛车子页 · 专注度驱动赛车界面", height_cm=10.0)

    add_heading(doc, "4.7 数字处方与设备配对（子页 / Prescription、Device）", 2)
    add_body(doc, (
        "数字处方子页根据 4 套处方包展示对应的章节列表与适用场景，章节支持单独播放进入呼吸训练；"
        "设备配对子页展示当前已连接的脑电头环型号、电池电量、采样率、固件版本，"
        "并实时呈现 4 个电极（Fp1、Fp2、TP9、TP10）的信号质量百分比，"
        "提供「检测电极接触 → 采集静息基线 → 校准情绪模型 → 完成」的 4 步校准流程入口。"
    ))
    add_figure_placeholder(doc, "图 4.7.1  数字处方子页 · 处方详情与章节列表", height_cm=8.5)
    add_figure_placeholder(doc, "图 4.7.2  设备配对子页 · 电极质量与校准流程", height_cm=8.5)


def chapter_5(doc: Document) -> None:
    add_heading(doc, "第 5 章  数据设计", 1)

    add_heading(doc, "5.1 核心数据类型", 2)
    add_body(doc, (
        "本系统的核心数据围绕「实时情绪状态快照」与「闭环干预记录」两条主线展开。"
        "下表列出主要 TypeScript 类型定义及其字段含义。"
    ))

    add_heading(doc, "5.1.1 ValenceSnapshot · 情绪效价快照", 3)
    add_table(
        doc,
        header=["字段", "类型", "取值范围", "含义"],
        rows=[
            ["id", "string", "v1 / v2 / …", "快照唯一标识"],
            ["time", "string", "HH:MM", "采样时刻"],
            ["score", "number", "0–100", "Valence 数值分数，越高越积极"],
            ["level", "ValenceLevel", "高效价 / 较高效价 / 较低效价 / 低效价", "效价等级（四分级）"],
            ["confidence", "number", "0–100", "模型置信度百分比"],
            ["signalQuality", "number", "0–100", "EEG 信号质量百分比"],
            ["eegSummary", "string", "—", "脑波节律自然语言摘要"],
            ["aiFeedback", "string", "—", "针对该快照预设的 AI 安抚语"],
            ["recommendedActionId", "string", "—", "推荐干预动作主键，关联 InterventionAction"],
        ],
    )

    add_heading(doc, "5.1.2 InterventionAction · 干预动作", 3)
    add_table(
        doc,
        header=["字段", "类型", "示例", "含义"],
        rows=[
            ["id", "string", "breath-alpha", "动作唯一标识"],
            ["title", "string", "3 分钟 Alpha 呼吸调节", "动作标题"],
            ["type", "枚举", "呼吸训练 / 数字处方 / 调节游戏 / AI陪伴 / 成长记录", "动作分类"],
            ["duration", "string", "3 分钟", "执行时长"],
            ["target", "string", "稳定较低效价状态", "目标 Valence 改善方向"],
            ["reason", "string", "—", "推荐理由文本"],
            ["route", "SubName", "player / race / aichat / prescription", "对应子页面路由名"],
            ["params", "Record<string,string>?", "{ id: 'p1' }", "跳转参数（可选）"],
        ],
    )

    add_heading(doc, "5.1.3 ClosedLoopRecord · 闭环记录", 3)
    add_table(
        doc,
        header=["字段", "类型", "含义"],
        rows=[
            ["id", "string", "记录唯一标识"],
            ["startTime", "string", "干预启动时间"],
            ["beforeLevel / beforeScore", "ValenceLevel / number", "干预前效价等级与分数"],
            ["action", "string", "本次执行的干预动作标题"],
            ["aiMessage", "string", "AI 在干预过程中给出的关键反馈语"],
            ["afterLevel / afterScore", "ValenceLevel / number", "干预后效价等级与分数"],
            ["delta", "number", "afterScore − beforeScore 增量"],
            ["status", "枚举", "已完成 / 进行中"],
        ],
    )

    add_heading(doc, "5.2 演示数据来源说明", 2)
    add_body(doc, (
        "为保证演示版本可独立运行且不涉及真实用户隐私数据，所有 Valence 快照、闭环记录、"
        "勋章、打卡进度等均为预置 mock 数据，集中维护在 [src/data/mock.ts](文件路径) 中。"
        "前端通过 useValenceLoop 钩子周期性轮换 Valence 快照池，并在每次轮换时合成少量随机抖动，"
        "模拟脑电信号的实时变化效果。后端不持久化任何 mock 数据。"
    ))


def chapter_6(doc: Document) -> None:
    add_heading(doc, "第 6 章  接口设计", 1)

    add_heading(doc, "6.1 AI 反馈接口  POST /api/chat", 2)
    add_table(
        doc,
        header=["属性", "说明"],
        rows=[
            ["接口路径", "POST /api/chat"],
            ["请求 Content-Type", "application/json; charset=utf-8"],
            ["响应 Content-Type", "application/json; charset=utf-8"],
            ["请求体", "{ messages: Array<{ role: 'system'|'user'|'assistant', content: string }> }"],
            ["请求约束", "messages 中至少包含 1 条 role=user；单条 content 长度上限 800 字符"],
            ["上下文注入", "前端会在 messages 最前面插入隐藏 system 消息携带 Valence 等级、score、置信度、信号质量、EEG 摘要"],
            ["响应体", "{ reply: string } 或 { error: string }"],
            ["响应状态码", "200 成功 / 400 参数错误 / 500 后端 / 502 上游空响应"],
        ],
    )

    add_heading(doc, "6.2 健康检查接口  GET /api/health", 2)
    add_table(
        doc,
        header=["属性", "说明"],
        rows=[
            ["接口路径", "GET /api/health"],
            ["响应 Content-Type", "application/json; charset=utf-8"],
            ["响应体", "{ ok: true, model: string, configured: boolean, keyFormatValid: boolean }"],
            ["用途", "前端可调用本接口确认后端代理是否已加载 DeepSeek API Key"],
            ["响应状态码", "200"],
        ],
    )

    add_heading(doc, "6.3 上游接口依赖", 2)
    add_body(doc, (
        "本系统后端将合并后的 messages 转发至 DeepSeek 官方 chat completions 接口，"
        "默认使用 deepseek-v4-flash 模型，温度参数 0.7，最大输出 token 数 600，非流式响应。"
        "上游接口地址、模型名称、API Key 均通过环境变量配置（DEEPSEEK_BASE_URL、"
        "DEEPSEEK_MODEL、DEEPSEEK_API_KEY），不在源码中硬编码。"
    ))


def chapter_7(doc: Document) -> None:
    add_heading(doc, "第 7 章  安全与合规", 1)

    add_heading(doc, "7.1 数据安全", 2)
    add_list_item(doc, "本系统演示版本不存储任何真实用户数据，所有交互结果仅保留在浏览器会话内；")
    add_list_item(doc, "EEG 信号与 Valence 推断结果均不上传第三方服务，仅在前端内存中维护；")
    add_list_item(doc, "DeepSeek API Key 由后端进程从环境变量加载，仅参与上游调用，不返回到前端；")
    add_list_item(doc, "后端对 messages 体长度做硬上限（最近 12 条），防止恶意构造长上下文；")
    add_list_item(doc, "对 API Key 进行 ASCII 字符合法性检测，避免误录入导致请求头注入异常。")

    add_heading(doc, "7.2 内容合规", 2)
    add_list_item(doc, "系统提示词明确要求 AI 角色不冒充医生、不做诊断、不给确定性医疗结论；")
    add_list_item(doc, "对于自伤、自杀、伤害他人等危机场景，AI 优先表达关切并引导用户联系线下专业资源；")
    add_list_item(doc, "全部 UI 文案禁止使用「治愈」「颠覆」「革命性」等夸张营销表达；")
    add_list_item(doc, "AI 陪伴页面同时提供「预约校心理中心」与「24h 援助热线」两个线下专业入口。")

    add_heading(doc, "7.3 兼容与可用性", 2)
    add_list_item(doc, "前端支持渐进式 Web 应用（PWA），可安装到桌面 / 移动设备主屏；")
    add_list_item(doc, "前端构建产物同时提供常规分块与单文件两种打包形态，便于离线分发；")
    add_list_item(doc, "界面采用响应式手机框布局，桌面与移动端视觉一致；")
    add_list_item(doc, "默认中文界面，技术术语以中英对照呈现（EEG、Valence、PWA、API 等）。")


def chapter_8(doc: Document) -> None:
    add_heading(doc, "第 8 章  运行结果与界面截图", 1)

    add_body(doc, (
        "下列截图按用户实际操作路径展示系统主要界面。运行环境：Windows 11、Chrome 130、"
        "Node.js 20.11、前端开发服务器 http://localhost:5173、后端 AI 代理 http://127.0.0.1:8787。"
    ))

    add_heading(doc, "8.1 启动与监测主页", 2)
    add_figure_placeholder(doc, "图 8.1.1  应用启动后的监测主页首屏", height_cm=11.0)

    add_heading(doc, "8.2 请求 AI 反馈", 2)
    add_figure_placeholder(doc, "图 8.2.1  点击「请求 AI 反馈」按钮后，AI 反馈卡刷新为 DeepSeek 真实回复", height_cm=11.0)

    add_heading(doc, "8.3 干预与训练", 2)
    add_figure_placeholder(doc, "图 8.3.1  干预页主屏 · 优先建议与可选干预动作", height_cm=11.0)

    add_heading(doc, "8.4 闭环记录与成长档案", 2)
    add_figure_placeholder(doc, "图 8.4.1  记录页 · 闭环时间轴与本月概览", height_cm=11.0)

    add_heading(doc, "8.5 AI 心理陪伴对话", 2)
    add_figure_placeholder(doc, "图 8.5.1  AI 陪伴子页 · 实时 Valence 状态条与多轮对话", height_cm=11.0)

    add_heading(doc, "8.6 呼吸训练播放器", 2)
    add_figure_placeholder(doc, "图 8.6.1  呼吸训练子页 · 4-7-8 节律光圈与放松度反馈", height_cm=11.0)

    add_heading(doc, "8.7 意念赛车调节游戏", 2)
    add_figure_placeholder(doc, "图 8.7.1  意念赛车子页 · 专注度驱动的赛车前进画面", height_cm=11.0)

    add_heading(doc, "8.8 数字处方与设备配对", 2)
    add_figure_placeholder(doc, "图 8.8.1  数字处方子页 · 章节列表与适用人群说明", height_cm=11.0)


def chapter_9(doc: Document) -> None:
    add_heading(doc, "第 9 章  后续可扩展方向", 1)

    add_body(doc, (
        "本版本（V1.0）已实现核心闭环主线与 8 个主要功能模块，可独立运行并完成完整的"
        "「监测 → 判断 → 反馈 → 干预 → 再监测」演示循环。后续版本规划在保留主线的前提下"
        "向以下方向扩展，扩展功能均明确标注为计划功能，不在本版本提交范围。"
    ))

    add_heading(doc, "9.1 真实硬件接入", 2)
    add_list_item(doc, "对接便携式干电极脑电头环（蓝牙 4.0+，256 Hz 采样），替换当前 mock 数据；")
    add_list_item(doc, "增加实时滤波（带通 1–40 Hz、陷波 50 Hz）与伪迹剔除模块；")
    add_list_item(doc, "Valence 推断从预置序列升级为基于 Alpha/Beta/Theta 频段功率比的实时模型。")

    add_heading(doc, "9.2 数据持久化", 2)
    add_list_item(doc, "本地 IndexedDB 存储用户的历史闭环记录、勋章进度、自定义干预偏好；")
    add_list_item(doc, "可选的云端同步（端到端加密），用户主动开启后才传输；")
    add_list_item(doc, "导出 PDF 周报，用户可分享给校园心理中心老师。")

    add_heading(doc, "9.3 评测与个性化", 2)
    add_list_item(doc, "干预效果统计分析（前后 Valence delta 分布、各干预手段的有效性排序）；")
    add_list_item(doc, "基于用户历史偏好的干预排序个性化（强化学习探索 / 利用平衡）；")
    add_list_item(doc, "校园心理中心面板（教师端）：以匿名群体维度查看趋势，不接触个人数据。")


def build_doc(out_path: Path) -> None:
    doc = Document()

    # 全局段落样式默认中文字体
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CN_FONT)

    # 第一节：A4 页面 + 页眉
    section = doc.sections[0]
    set_a4_margins(section)
    add_header(section)

    # 封面
    add_cover(doc)

    # 各章节
    chapter_1(doc); add_page_break(doc)
    chapter_2(doc); add_page_break(doc)
    chapter_3(doc); add_page_break(doc)
    chapter_4(doc); add_page_break(doc)
    chapter_5(doc); add_page_break(doc)
    chapter_6(doc); add_page_break(doc)
    chapter_7(doc); add_page_break(doc)
    chapter_8(doc); add_page_break(doc)
    chapter_9(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[完成] 输出：{out_path}")


if __name__ == "__main__":
    build_doc(OUT)
